import docx
import os

base = r"c:\Users\SHERI VAIJAYANTH\Desktop\Projects\Open Energy Grid Explorer"

for fname in ["PRD.docx", "SYSTEM ARCHITECTURE.docx"]:
    path = os.path.join(base, fname)
    out_path = os.path.join(base, fname.replace(".docx", ".txt"))
    doc = docx.Document(path)
    
    lines = []
    for p in doc.paragraphs:
        lines.append(p.text)
    
    # Also extract tables
    for i, table in enumerate(doc.tables):
        lines.append(f"\n--- TABLE {i+1} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    
    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(f"Extracted {fname} -> {out_path}")
